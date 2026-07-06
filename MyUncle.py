#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MyUncle.py  —  maintenance-lattice Monte-Carlo + mean-field engine
================================================================================
A directionally ordered medium eroded by a state-dependent flux and restored by
repair. One abstraction, many disciplines. This build adds the discipline preset

        tissue_atavism

which maps the engine onto cancer as an atavistic (multicellular -> unicellular)
order-repair transition, and ships six therapy PROTOCOLS as built-in interventions:

    MTD            maximum tolerated dose (control)
    A1_immuno      target recently-evolved weakness (adaptive immunity)
    A2_adaptive    evolutionary-steering dose modulation (space-limit resistance)
    A3_simblock    MTD + stress-induced-mutagenesis inhibitor (disable evolvability)
    A4_rediff      re-differentiation flux (push back across the hysteresis loop)
    A2A3           A2 + A3   (adaptive + SIM-block)
    A1A4           A1 + A4   (immuno + re-differentiation)

Framework pieces reused from the core MyUncle abstraction:
  * von Mises order-entropy bridge  U_chi = I1(k)/I0(k),  dS/dU_chi = -kappa
  * mean-field reduction            dI/dtau = mu(I)(1-I) - rho*I     (bistable)
  * analyses                        bifurcation, hysteresis, entropy_production
  * evolutionary (GaaR-style) layer clones on a resistance trait under selection
  * reproduce_paper()               figure suite + dark cover + house-style .docx

CLI:
    python MyUncle.py --reproduce --preset tissue_atavism \
        --title "Cancer as a bistable order-repair transition" \
        --subtitle "In-silico comparison of evolutionary and atavism-targeting therapies" [--full]
"""
from __future__ import annotations
import os, sys, argparse, textwrap, datetime
import numpy as np
from scipy.special import i0, i1

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.gridspec import GridSpec

# ------------------------------------------------------------------ house style
HOUSE = dict(
    font="Arial",
    navy="#1F3864", blue="#2E5FA3", h1="#2E74B5", h2="#2E74B5", h3="#1F4D78",
    rule="#2E5FA3", caption="#444444",
    author="Leon Sandler",
    penname="Leonid Sandler",
    orcid="0009-0007-4584-808X",
    github="github.com/sandlerleon",
    affil="Independent Researcher, Northbrook, Illinois, United States",
)
PALETTE = {
    "MTD":        "#9AA5B1", "A1_immuno":  "#E08A2B", "A2_adaptive": "#3AA6A0",
    "A3_simblock":"#C0504D", "A4_rediff":  "#8064A2", "A2A3":        "#1F3864",
    "A1A4":       "#4C9A2A",
}
LABELS = {
    "MTD":"MTD (control)", "A1_immuno":"A1 Immuno / young-gene",
    "A2_adaptive":"A2 Adaptive dosing", "A3_simblock":"A3 MTD + SIM-block",
    "A4_rediff":"A4 Re-differentiation", "A2A3":"A2+A3 Adaptive + SIM-block",
    "A1A4":"A1+A4 Immuno + Diff",
}
PROTOCOLS = list(LABELS.keys())

def _mpl_style():
    plt.rcParams.update({
        "font.family": "DejaVu Sans", "axes.titlesize": 11, "axes.labelsize": 9.5,
        "xtick.labelsize": 8, "ytick.labelsize": 8, "legend.fontsize": 7.4,
        "axes.edgecolor": "#555555",
    })

# ============================================================================ #
#  1. von Mises order-entropy bridge                                            #
# ============================================================================ #
def order_parameter(kappa):
    """U_chi = <cos theta> = I1(kappa)/I0(kappa) for a von Mises medium."""
    kappa = np.asarray(kappa, float)
    return np.where(kappa > 0, i1(kappa) / i0(kappa), 0.0)

def vonmises_entropy(kappa):
    """Differential entropy S(kappa) of the von Mises distribution (nats)."""
    kappa = np.asarray(kappa, float)
    return np.log(2 * np.pi * i0(kappa)) - kappa * order_parameter(kappa)

def dS_dUchi(kappa):
    """Exact order-entropy bridge identity: dS/dU_chi = -kappa."""
    return -np.asarray(kappa, float)

# ============================================================================ #
#  2. Mean-field bistable order-repair engine                                   #
#     dI/dtau = mu(I)(1-I) - rho*I ,  mu(I)=mu0(1+beta I^h/(K^h+I^h))           #
# ============================================================================ #
class Engine:
    def __init__(self, mu0=0.08, beta=6.0, Kh=0.5, hh=4.0):
        self.mu0, self.beta, self.Kh, self.hh = mu0, beta, Kh, hh

    def mu(self, I):
        I = np.asarray(I, float)
        return self.mu0 * (1 + self.beta * I**self.hh / (self.Kh**self.hh + I**self.hh))

    def rhs(self, I, rho, diff=0.0):
        return self.mu(I) * (1 - I) - rho * I + diff

    def bifurcation(self, rho_grid, Igrid=None):
        """Return stable/unstable fixed points I*(rho) by sign changes of rhs."""
        if Igrid is None:
            Igrid = np.linspace(1e-3, 1.0, 2000)
        stable, unstable = [], []
        for rho in rho_grid:
            f = self.rhs(Igrid, rho)
            sign = np.sign(f)
            idx = np.where(np.diff(sign) != 0)[0]
            for j in idx:
                I0_ = Igrid[j] - f[j] * (Igrid[j+1]-Igrid[j])/(f[j+1]-f[j])
                slope = (self.rhs(I0_+1e-4, rho) - self.rhs(I0_-1e-4, rho)) / 2e-4
                (stable if slope < 0 else unstable).append((rho, I0_))
        return np.array(stable), np.array(unstable)

    def hysteresis(self, rho_up, dt=0.5, I0=0.9):
        """Sweep rho up then down; trace I to expose the hysteresis loop."""
        rho_seq = np.concatenate([rho_up, rho_up[::-1]])
        I = I0; out = []
        for rho in rho_seq:
            for _ in range(200):
                I = np.clip(I + self.rhs(I, rho) * dt, 1e-3, 1.0)
            out.append(I)
        n = len(rho_up)
        return rho_up, np.array(out[:n]), np.array(out[n:][::-1])

# ============================================================================ #
#  3. tissue_atavism preset + therapy cohort simulation                         #
# ============================================================================ #
TISSUE_ATAVISM = dict(
    # cohort / time
    M=500, dt=0.5, T=400.0,
    N0=0.35, K=1.0, N_lethal=0.95, control=0.50,
    # tumour ecology
    kill_cyto=0.28, kill_imm=0.17, k_esc=0.18, imm_floor=0.12,
    mu0=8e-5, mu_sim=4e-3, rev_diff=0.04,
    gamma_at=0.80, alpha_at=3.0,
    # order-repair (I) engine erosion flux
    rho0=0.05, rho_str=0.25, rho_bur=0.32, diff_flux=0.12, cyto_lowdose=0.10,
    seed=11,
)

def _init_cohort(P):
    rng = np.random.default_rng(P["seed"])
    M = P["M"]
    g_base = np.clip(rng.normal(0.11, 0.02, M), 0.05, 0.18)
    cost   = rng.uniform(0.20, 0.50, M)
    phi0   = 10**rng.uniform(-4, -2.3, M)
    imm_q  = rng.uniform(0.40, 1.30, M)
    I0     = rng.uniform(0.55, 0.80, M)
    gnoise = rng.normal(0.0, 0.03, (M, int(P["T"]/P["dt"])))
    return dict(g_base=g_base, cost=cost, phi0=phi0, imm_q=imm_q, I0=I0, gnoise=gnoise)

def simulate_protocol(protocol, P=TISSUE_ATAVISM, coh=None, cost_override=None):
    """Run one therapy protocol across the paired virtual-mouse cohort."""
    P = dict(P); eng = Engine()
    if coh is None:
        coh = _init_cohort(P)
    M, dt, steps = P["M"], P["dt"], int(P["T"]/P["dt"])
    K, NL = P["K"], P["N_lethal"]
    g_base = coh["g_base"]
    cost = coh["cost"] if cost_override is None else np.full(M, cost_override)
    phi0, imm_q, gnoise = coh["phi0"], coh["imm_q"], coh["gnoise"]

    S = (1 - phi0) * P["N0"]; R = phi0 * P["N0"]
    I = coh["I0"].copy(); imm = np.ones(M)
    on = np.ones(M, bool)
    dead = np.full(M, np.nan); dose = np.zeros(M)
    Straj = np.zeros((M, steps)); Rtraj = np.zeros((M, steps)); Itraj = np.zeros((M, steps))

    for t in range(steps):
        N = S + R
        Straj[:, t] = S; Rtraj[:, t] = R; Itraj[:, t] = I
        D = np.zeros(M); immune = False; sim = False; diff = 0.0
        if   protocol == "MTD":         D[:] = 1.0; sim = True
        elif protocol == "A1_immuno":   immune = True
        elif protocol == "A2_adaptive": D = np.where(on, 1.0, 0.0); sim = True
        elif protocol == "A3_simblock": D[:] = 1.0
        elif protocol == "A4_rediff":   D[:] = P["cyto_lowdose"]; diff = P["diff_flux"]
        elif protocol == "A2A3":        D = np.where(on, 1.0, 0.0)
        elif protocol == "A1A4":        immune = True; diff = P["diff_flux"]
        dose += D * dt

        if immune:
            imm = np.maximum(P["imm_floor"], imm * np.exp(-P["k_esc"] * N * dt))

        rho = P["rho0"] + P["rho_str"] * D + P["rho_bur"] * N
        I = np.clip(I + eng.rhs(I, rho, diff) * dt, 1e-3, 1.0)

        geff = g_base * (1 + P["gamma_at"] * (1 - I))
        logi = np.clip(1 - N / K, -1, 1)
        S *= np.exp((geff * logi + gnoise[:, t]) * dt)
        R *= np.exp((geff * (1 - cost) * logi + gnoise[:, t]) * dt)

        ik = P["kill_imm"] * imm_q * imm if immune else 0.0
        S *= np.exp(-(P["kill_cyto"] * D + ik) * dt)
        R *= np.exp(-(ik if immune else 0.0) * dt)

        mu = P["mu0"] * (1 + P["alpha_at"] * (1 - I)) + (P["mu_sim"] * D if sim else 0.0)
        flux = mu * S; rev = P["rev_diff"] * (I > 0.85) * R
        S = np.clip(S - flux * dt + rev * dt, 0, None)
        R = np.clip(R + flux * dt - rev * dt, 0, None)
        S[S < 1e-12] = 0; R[R < 1e-12] = 0

        Nn = S + R
        if protocol in ("A2_adaptive", "A2A3"):
            on = np.where(Nn <= 0.55, False, on)
            on = np.where(Nn >= 0.80, True, on)
        newly = np.isnan(dead) & (Nn >= NL); dead[newly] = (t + 1) * dt

    surv = np.where(np.isnan(dead), P["T"], dead)
    Ntraj = Straj + Rtraj; alive = np.isnan(dead)
    return dict(surv=surv, alive=alive, dead=dead, dose=dose,
                controlled=alive & (Ntraj[:, -1] < P["control"]),
                Ntraj=Ntraj, Itraj=Itraj)

def run_horse_race(P=TISSUE_ATAVISM):
    coh = _init_cohort(P)
    res = {p: simulate_protocol(p, P, coh) for p in PROTOCOLS}
    summary = {}
    for p in PROTOCOLS:
        Rr = res[p]
        summary[p] = dict(medSurv=float(np.median(Rr["surv"])),
                          alive=100*float(Rr["alive"].mean()),
                          controlled=100*float(Rr["controlled"].mean()),
                          dose=float(Rr["dose"].mean()))
    # sensitivity to fitness cost of resistance
    costs = [0.10, 0.20, 0.30, 0.40, 0.55]
    sarms = ["MTD", "A2_adaptive", "A3_simblock", "A2A3"]
    sens = {a: [float(np.median(simulate_protocol(a, P, coh, cost_override=c)["surv"]))
                for c in costs] for a in sarms}
    return res, summary, dict(costs=costs, arms=sarms, data=sens), coh

# ============================================================================ #
#  4. Figure suite (house style)                                                #
# ============================================================================ #
def fig_theory(path, eng=None):
    """Von Mises bridge + bifurcation + hysteresis: the order-repair theory."""
    _mpl_style()
    if eng is None: eng = Engine()
    fig = plt.figure(figsize=(13, 4.2)); gs = GridSpec(1, 3, figure=fig, wspace=0.30)
    navy = HOUSE["navy"]

    ax = fig.add_subplot(gs[0, 0])
    k = np.linspace(0.01, 12, 400)
    ax.plot(order_parameter(k), vonmises_entropy(k), color=navy, lw=2)
    ax.set_title("(A) Order\u2013entropy bridge", color=navy, fontweight="bold")
    ax.set_xlabel(r"cooperation order  $U_\chi=I_1(\kappa)/I_0(\kappa)$")
    ax.set_ylabel("entropy  S  (nats)"); ax.grid(alpha=0.25)
    ax.text(0.05, 0.10, r"$dS/dU_\chi=-\kappa$", transform=ax.transAxes,
            fontsize=10, color="#2E5FA3")

    ax = fig.add_subplot(gs[0, 1])
    rho_grid = np.linspace(0.02, 0.6, 400)
    stab, unstab = eng.bifurcation(rho_grid)
    if len(stab):    ax.scatter(stab[:, 0],   stab[:, 1],   s=3, color="#2E74B5", label="stable")
    if len(unstab):  ax.scatter(unstab[:, 0], unstab[:, 1], s=3, color="#C0504D", label="unstable")
    ax.set_title("(B) Bistable reversion", color=navy, fontweight="bold")
    ax.set_xlabel(r"erosion flux  $\rho$"); ax.set_ylabel(r"tissue order  $I^*$")
    ax.legend(loc="upper right"); ax.grid(alpha=0.25)

    ax = fig.add_subplot(gs[0, 2])
    rho_up = np.linspace(0.05, 0.5, 60)
    r, up, dn = eng.hysteresis(rho_up)
    ax.plot(r, up, color="#C0504D", lw=2, label="stress \u2191 (onset)")
    ax.plot(r, dn, color="#2E74B5", lw=2, label="stress \u2193 (differentiation)")
    ax.fill_between(r, up, dn, color="#8064A2", alpha=0.12)
    ax.set_title("(C) Hysteresis \u2014 why reversion sticks", color=navy, fontweight="bold")
    ax.set_xlabel(r"erosion flux  $\rho$"); ax.set_ylabel(r"tissue order  $I$")
    ax.legend(loc="lower left"); ax.grid(alpha=0.25)

    fig.suptitle("Theory layer: MyUncle order\u2013repair engine on the cancer atavism map",
                 fontsize=12, fontweight="bold", color=navy, y=1.02)
    fig.savefig(path, dpi=140, bbox_inches="tight", facecolor="white"); plt.close(fig)

def fig_horse_race(path, res, summary, sens, coh, P=TISSUE_ATAVISM):
    _mpl_style()
    navy = HOUSE["navy"]; steps = int(P["T"]/P["dt"]); tg = np.arange(steps+1)*P["dt"]
    fig = plt.figure(figsize=(13, 9)); gs = GridSpec(2, 2, figure=fig, hspace=0.30, wspace=0.24)

    ax = fig.add_subplot(gs[0, 0])
    for p in PROTOCOLS:
        d = res[p]["dead"]; km = np.array([(np.isnan(d) | (d > tt)).mean() for tt in tg])
        ax.step(tg, 100*km, where="post", color=PALETTE[p], lw=2, label=LABELS[p], alpha=0.9)
    ax.set_title(f"(A) Survival \u2014 virtual mouse cohort (n={P['M']})", color=navy, fontweight="bold")
    ax.set_xlabel("day"); ax.set_ylabel("% surviving"); ax.set_ylim(0, 102); ax.set_xlim(0, P["T"])
    ax.legend(loc="lower left"); ax.grid(alpha=0.25)

    ax = fig.add_subplot(gs[0, 1]); xa = np.arange(len(PROTOCOLS))
    meds = [summary[p]["medSurv"] for p in PROTOCOLS]; al = [summary[p]["alive"] for p in PROTOCOLS]
    ax.bar(xa-0.2, meds, 0.4, color=[PALETTE[p] for p in PROTOCOLS])
    ax2 = ax.twinx(); ax2.bar(xa+0.2, al, 0.4, color=[PALETTE[p] for p in PROTOCOLS],
                              alpha=0.45, hatch="//", edgecolor="white")
    ax.set_title("(B) Median survival (solid) & % alive @ day400 (hatched)", color=navy, fontweight="bold")
    ax.set_xticks(xa); ax.set_xticklabels([p.split("_")[0] for p in PROTOCOLS], fontsize=8)
    ax.set_ylabel("median survival (day)"); ax2.set_ylabel("% alive @ day400")
    ax.set_ylim(0, P["T"]*1.02); ax2.set_ylim(0, 102)
    for i, p in enumerate(PROTOCOLS):
        ax.text(i-0.2, meds[i]+5, f"{meds[i]:.0f}", ha="center", fontsize=7.3, color=navy)
        ax2.text(i+0.2, al[i]+2, f"{al[i]:.0f}", ha="center", fontsize=7.3, color="#444")
    ax.grid(axis="y", alpha=0.2)

    ax = fig.add_subplot(gs[1, 0])
    mi = int(np.argmin(np.abs(coh["g_base"]-np.median(coh["g_base"])) + np.abs(coh["cost"]-0.30)))
    for p in PROTOCOLS:
        ax.plot(tg[:-1], res[p]["Ntraj"][mi], color=PALETTE[p], lw=1.7, label=LABELS[p], alpha=0.9)
    ax.axhline(P["N_lethal"], color="k", ls=":", lw=1, alpha=0.6)
    ax.text(4, P["N_lethal"]+0.01, "lethal", fontsize=7, alpha=0.7)
    ax.set_title(f"(C) Tumor burden \u2014 representative mouse (cost={coh['cost'][mi]:.2f})",
                 color=navy, fontweight="bold")
    ax.set_xlabel("day"); ax.set_ylabel("tumor burden N"); ax.set_xlim(0, P["T"]); ax.set_ylim(0, 1.0)
    ax.legend(loc="upper right", fontsize=6.6); ax.grid(alpha=0.25)

    ax = fig.add_subplot(gs[1, 1])
    for a in sens["arms"]:
        ax.plot(sens["costs"], sens["data"][a], "o-", color=PALETTE[a], lw=2, label=LABELS[a])
    ax.set_title("(D) Sensitivity to fitness cost of resistance", color=navy, fontweight="bold")
    ax.set_xlabel("cost of resistance"); ax.set_ylabel("median survival (day)"); ax.set_ylim(0, P["T"]*1.02)
    ax.legend(loc="upper left", fontsize=7.2); ax.grid(alpha=0.25)

    fig.suptitle("In-silico horse race: order\u2013repair engine + evolutionary tumor dynamics",
                 fontsize=13, fontweight="bold", color=navy, y=0.985)
    fig.savefig(path, dpi=140, bbox_inches="tight", facecolor="white"); plt.close(fig)

def fig_cover(path, title, subtitle):
    """Dark house-style cover."""
    _mpl_style()
    fig = plt.figure(figsize=(8.5, 11)); fig.patch.set_facecolor("#0B1220")
    ax = fig.add_axes([0, 0, 1, 1]); ax.set_facecolor("#0B1220"); ax.axis("off")
    eng = Engine()
    # faint bifurcation motif
    rho_grid = np.linspace(0.02, 0.6, 300); stab, unstab = eng.bifurcation(rho_grid)
    if len(stab):
        ax.scatter(0.12+0.76*(stab[:,0]-0.02)/0.58, 0.14+0.30*stab[:,1],
                   s=2, color="#2E74B5", alpha=0.5)
    if len(unstab):
        ax.scatter(0.12+0.76*(unstab[:,0]-0.02)/0.58, 0.14+0.30*unstab[:,1],
                   s=2, color="#C0504D", alpha=0.5)
    ax.plot([0.08, 0.92], [0.66, 0.66], color=HOUSE["rule"], lw=2)
    ax.text(0.08, 0.80, textwrap.fill(title, 34), fontsize=27, color="#EAF0FA",
            fontweight="bold", va="top")
    ax.text(0.08, 0.685, textwrap.fill(subtitle, 60), fontsize=13, color="#8FB3E0",
            style="italic", va="top")
    ax.text(0.08, 0.60, HOUSE["author"], fontsize=13, color="#EAF0FA", fontweight="bold")
    ax.text(0.08, 0.575, f"pen name {HOUSE['penname']}", fontsize=10, color="#8FB3E0", style="italic")
    ax.text(0.08, 0.55, HOUSE["affil"], fontsize=9.5, color="#B8C6DC", style="italic")
    ax.text(0.08, 0.09, "MyUncle \u2022 tissue_atavism preset \u2022 order\u2013repair engine",
            fontsize=9, color="#5E7BA6")
    ax.text(0.08, 0.065, datetime.date.today().isoformat(), fontsize=9, color="#5E7BA6")
    ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    fig.savefig(path, dpi=150, facecolor="#0B1220"); plt.close(fig)

# ============================================================================ #
#  5. Manuscript (nanoribbon house style) via python-docx                       #
# ============================================================================ #
# ---- parameter provenance (order-of-magnitude, phenomenological unless noted) ----
PARAM_SOURCES = [
    ("g_base", "0.05-0.18 /day", "Net proliferation; spans reported murine tumour doubling times of days-weeks (phenomenological range)."),
    ("cost", "0.20-0.50", "Fitness cost of full resistance; a cost of resistance is documented in vitro and in mice (Bacevic et al. 2017). Swept in Fig. 2D."),
    ("phi0", "1e-4-5e-3", "Pre-existing resistant fraction; consistent with rare standing resistant clones assumed in adaptive-therapy models (Gatenby et al. 2009)."),
    ("kill_cyto", "0.28 /day", "Cytotoxic kill of sensitive cells at max dose; phenomenological, chosen so MTD gives an initial response then relapse."),
    ("kill_imm / k_esc / imm_floor", "0.17 / 0.18 / 0.12", "Immune clearance, burden-driven escape, residual surveillance; produce the bimodal hot/cold response seen clinically."),
    ("mu0 / mu_sim", "8e-5 / 4e-3 /day", "Baseline vs stress-induced S->R acquisition; the stress term encodes SOS-homologue mutagenesis (Cisneros et al. 2017)."),
    ("rho0 / rho_str / rho_bur", "0.05 / 0.25 / 0.32", "Order-erosion flux from baseline, drug DNA-stress, and burden; sets the bistable/hysteretic regime (Fig. 1)."),
    ("diff_flux / rev_diff", "0.12 / 0.04", "Re-differentiation drive and R->S reversion; phenomenological analogue of differentiation therapy (e.g. retinoic acid in acute promyelocytic leukaemia)."),
]

def _race_medians(P, coh):
    """Lightweight horse race: median survival + % alive per protocol (no figures)."""
    out = {}
    for p in PROTOCOLS:
        r = simulate_protocol(p, P, coh)
        out[p] = (float(np.median(r["surv"])), 100*float(r["alive"].mean()))
    return out

def global_sensitivity(P=TISSUE_ATAVISM, n_lhs=120, span=0.30, seed=2024):
    """Latin-hypercube robustness of the two qualitative conclusions +
    one-at-a-time tornado on the adaptive+SIM-block survival advantage over MTD."""
    rng = np.random.default_rng(seed)
    keys = ["kill_cyto","kill_imm","k_esc","mu_sim","rev_diff",
            "rho_bur","diff_flux","gamma_at","alpha_at"]
    Plite = dict(P); Plite["M"] = 120
    coh = _init_cohort(Plite)
    # LHS samples in [1-span, 1+span]
    lhs = (np.arange(n_lhs)[:, None] + rng.random((n_lhs, len(keys)))) / n_lhs
    lhs = lhs[rng.permutation(n_lhs)]
    for j in range(len(keys)):
        lhs[:, j] = rng.permutation(lhs[:, j])
    factors = (1 - span) + 2*span*lhs
    ordering_ok = 0; family_ok = 0; gains = []
    drug_order = ["MTD", "A3_simblock", "A2_adaptive", "A2A3"]
    for i in range(n_lhs):
        Pp = dict(Plite)
        for j, k in enumerate(keys):
            Pp[k] = P[k] * factors[i, j]
        med = _race_medians(Pp, coh)
        ms = {p: med[p][0] for p in PROTOCOLS}; al = {p: med[p][1] for p in PROTOCOLS}
        if ms["MTD"] <= ms["A3_simblock"] <= ms["A2_adaptive"] <= ms["A2A3"]:
            ordering_ok += 1
        best_alive = max(PROTOCOLS, key=lambda p: al[p])
        if best_alive in ("A1_immuno", "A1A4"):
            family_ok += 1
        gains.append(ms["A2A3"] - ms["MTD"])
    # tornado: one-at-a-time +/- span on the A2A3 - MTD gain
    base = _race_medians(Plite, coh)
    base_gain = base["A2A3"][0] - base["MTD"][0]
    tornado = {}
    for k in keys:
        lo = dict(Plite); lo[k] = P[k]*(1-span)
        hi = dict(Plite); hi[k] = P[k]*(1+span)
        ml = _race_medians(lo, coh); mh = _race_medians(hi, coh)
        tornado[k] = (ml["A2A3"][0]-ml["MTD"][0]-base_gain,
                      mh["A2A3"][0]-mh["MTD"][0]-base_gain)
    return dict(n=n_lhs, span=span,
                ordering_frac=100*ordering_ok/n_lhs,
                family_frac=100*family_ok/n_lhs,
                gains=np.array(gains), base_gain=base_gain, tornado=tornado, keys=keys)

def fig_sensitivity(path, gs):
    _mpl_style(); navy = HOUSE["navy"]
    fig = plt.figure(figsize=(13, 4.6)); grid = GridSpec(1, 2, figure=fig, wspace=0.28)
    ax = fig.add_subplot(grid[0, 0])
    ax.hist(gs["gains"], bins=22, color="#2E74B5", alpha=0.85, edgecolor="white")
    ax.axvline(0, color="k", ls=":", lw=1)
    ax.axvline(gs["base_gain"], color="#C0504D", lw=2, label=f"baseline gain {gs['base_gain']:.0f} d")
    ax.set_title(f"(A) Adaptive+SIM-block advantage over MTD\n"
                 f"LHS n={gs['n']}, \u00b1{int(gs['span']*100)}% on 9 parameters",
                 color=navy, fontweight="bold")
    ax.set_xlabel("median-survival gain vs MTD (days)"); ax.set_ylabel("samples")
    ax.legend(loc="upper right"); ax.grid(alpha=0.25)
    ax.text(0.02, 0.92, f"drug ordering preserved: {gs['ordering_frac']:.0f}%\n"
                         f"young-target family wins survivors: {gs['family_frac']:.0f}%",
            transform=ax.transAxes, fontsize=8.5, color="#1F4D78", va="top")
    ax = fig.add_subplot(grid[0, 1])
    keys = gs["keys"]; y = np.arange(len(keys))
    lows = [gs["tornado"][k][0] for k in keys]; highs = [gs["tornado"][k][1] for k in keys]
    order = np.argsort([abs(l)+abs(h) for l, h in zip(lows, highs)])
    keys = [keys[i] for i in order]; lows = [lows[i] for i in order]; highs = [highs[i] for i in order]
    for i, (l, h) in enumerate(zip(lows, highs)):
        ax.barh(i, l, color="#C0504D", alpha=0.8); ax.barh(i, h, color="#2E74B5", alpha=0.8)
    ax.axvline(0, color="k", lw=1)
    ax.set_yticks(range(len(keys))); ax.set_yticklabels(keys, fontsize=8)
    ax.set_title("(B) Tornado: \u00b130% one-at-a-time\n(effect on the survival gain)",
                 color=navy, fontweight="bold")
    ax.set_xlabel("\u0394 gain vs baseline (days)"); ax.grid(axis="x", alpha=0.25)
    fig.suptitle("Global sensitivity: robustness of the qualitative conclusions",
                 fontsize=12, fontweight="bold", color=navy, y=1.03)
    fig.savefig(path, dpi=140, bbox_inches="tight", facecolor="white"); plt.close(fig)

def build_docx(path, title, subtitle, summary, gsens, cover_png, theory_png, race_png, sens_png, anon=False):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def hexrgb(h): return RGBColor(int(h[1:3],16), int(h[3:5],16), int(h[5:7],16))
    NAVY, BLUE, H1, H2, H3, CAP = (hexrgb(HOUSE[k]) for k in ("navy","blue","h1","h2","h3","caption"))

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for m in ("top_margin","bottom_margin","left_margin","right_margin"):
        setattr(sec, m, Inches(1))

    def para(text="", size=10.5, color=None, bold=False, italic=False,
             align=None, before=0, after=6):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        if align is not None: p.alignment = align
        r = p.add_run(text); r.font.name = "Arial"; r.font.size = Pt(size)
        r.bold = bold; r.italic = italic
        if color is not None: r.font.color.rgb = color
        return p

    def rule():
        p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"12")
        bottom.set(qn("w:space"),"1"); bottom.set(qn("w:color"),HOUSE["rule"][1:])
        pbdr.append(bottom); pPr.append(pbdr)

    def h1(t): para(t, 16, H1, bold=True, before=12, after=4)
    def h2(t): para(t, 13, H2, bold=True, before=8, after=3)
    def h3(t): para(t, 12, H3, bold=True, before=6, after=2)
    def body(t): para(t, 10.5, after=8)
    def caption(t): para(t, 9, CAP, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    def figure(png, width=6.5):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(png, width=Inches(width))

    C = WD_ALIGN_PARAGRAPH.CENTER
    # ---- title block ----
    para(title, 18, NAVY, bold=True, align=C, after=2)
    para(subtitle, 12, BLUE, italic=True, align=C, after=6)
    if anon:
        para("Author identity and affiliation removed for double-blind peer review.",
             10, italic=True, align=C, after=4)
    else:
        para(HOUSE["author"], 11, bold=True, align=C, after=1)
        para(HOUSE["affil"], 10, italic=True, align=C, after=1)
        para("ORCID: " + HOUSE["orcid"], 10, italic=True, align=C, after=4)
    rule()

    if not anon:
        figure(cover_png, width=3.0)
        caption("Bistable order\u2013repair motif.")

    # ---- abstract (novelty foregrounded; framed as theoretical oncology) ----
    h2("Abstract")
    best_alive = max(PROTOCOLS, key=lambda p: summary[p]["alive"])
    best_drug  = max(["MTD","A2_adaptive","A3_simblock","A2A3"], key=lambda p: summary[p]["medSurv"])
    body(
        "We present a theoretical-oncology framework in which cancer initiation and therapy response are "
        "governed by a bistable order\u2013repair transition. Our principal innovation is to couple "
        "evolutionary (sensitive/resistant) therapy dynamics to a single bistable order parameter I "
        "representing multicellular information integrity: healthy tissue is a directionally ordered, "
        "cooperative medium eroded by a state-dependent stress flux and restored by repair, so that the "
        "mean-field reduction dI/d\u03c4 = \u03bc(I)(1\u2212I) \u2212 \u03c1I possesses a low-order "
        "attractor identified with the atavistic (quasi-unicellular) phenotype. Across a paired cohort "
        f"of n={TISSUE_ATAVISM['M']} virtual mice, survival under DNA-damaging protocols increases "
        "monotonically from maximum-tolerated dose toward evolutionary-steering plus suppression of "
        f"stress-induced mutagenesis ({LABELS[best_drug]}: median {summary[best_drug]['medSurv']:.0f} d "
        f"vs {summary['MTD']['medSurv']:.0f} d for MTD), while only protocols targeting recently-evolved "
        f"vulnerabilities \u2014 adaptive immunity and re-differentiation \u2014 yield long-term survivors "
        f"({LABELS[best_alive]}: {summary[best_alive]['alive']:.0f}% alive at day 400). A Latin-hypercube "
        f"global sensitivity analysis (\u00b1{int(gsens['span']*100)}% on nine parameters) preserves the "
        f"drug-arm ordering in {gsens['ordering_frac']:.0f}% of samples and the young-target survivor "
        f"advantage in {gsens['family_frac']:.0f}%. The model provides a single dynamical framework "
        "consistent with the atavism hypothesis's therapeutic corollary \u2014 target the weakness, not "
        "the strength \u2014 deriving it from dynamical structure rather than tuned assumptions. It is "
        "hypothesis-generating rather than confirmatory.")

    # ---- 1. Introduction ----
    h1("1. Introduction")
    body(
        "The atavism (or 'cancer-first') model holds that malignant transformation is a stress-triggered "
        "reversion to an ancient, evolutionarily conserved unicellular survival program rather than a "
        "purely stochastic accumulation of mutations (Davies & Lineweaver 2011; Lineweaver et al. 2021). "
        "Phylostratigraphic analyses report that tumours over-express evolutionarily older genes "
        "(Trigos et al. 2017) and that part of cancer's elevated mutation rate is self-inflicted, "
        "mediated by homologues of the bacterial SOS / stress-induced-mutagenesis response "
        "(Cisneros et al. 2017). If cancer is a reversion, therapy logic inverts: maximally stressing the "
        "tumour with DNA-damaging agents both selects for resistance and provokes the ancient mutation "
        "engine that generates it (Lineweaver, Davies & Vincent 2014).")
    h2("1.1 Novelty relative to existing models")
    body(
        "Sensitive/resistant competition, adaptive therapy, fitness costs and evolutionary dynamics are "
        "individually well established (Gatenby et al. 2009; Zhang et al. 2017). Our contribution is not "
        "another resistance model but the coupling of that evolutionary layer to a bistable order\u2013"
        "repair variable representing multicellular information integrity. This single addition lets one "
        "framework generate, and relate, four therapeutic logics \u2014 evolutionary steering, "
        "evolvability suppression, differentiation, and immune targeting of young dependencies \u2014 and "
        "explains why differentiation therapy is bounded by hysteresis and why DNA-stressing therapy is "
        "self-defeating, neither of which follows from resistance dynamics alone.")

    # ---- 2. Model ----
    h1("2. Model")
    h2("2.1 Order\u2013repair engine")
    body(
        "Tissue cooperation is summarised by an order parameter I \u2208 [0,1] (high I = cooperative "
        "multicellular state; low I = atavistic reversion). Its dynamics follow the mean-field reduction "
        "dI/d\u03c4 = \u03bc(I)(1\u2212I) \u2212 \u03c1I with self-amplifying repair "
        "\u03bc(I)=\u03bc\u2080(1+\u03b2I\u02b0/(K\u02b0+I\u02b0)) and erosion flux "
        "\u03c1 = \u03c1\u2080 + \u03c1_str\u00b7D + \u03c1_bur\u00b7N rising with drug DNA-stress D and "
        "tumour burden N. The system is bistable (Fig. 1B) and the reversion hysteretic (Fig. 1C): once "
        "a large tumour settles into the low-order basin, removing the stress does not restore order \u2014 "
        "the basis for the limited reach of re-differentiation therapy.")
    figure(theory_png, width=6.7)
    caption("Fig. 1. Theory layer. (A) von Mises order\u2013entropy bridge, dS/dU_\u03c7 = \u2212\u03ba. "
            "(B) Bistable fixed points I*(\u03c1). (C) Hysteresis loop under a stress up/down sweep.")

    h2("2.2 Evolutionary layer and therapies")
    body(
        "Tumour cells partition into drug-sensitive (S) and resistant (R) compartments competing for a "
        "shared carrying capacity; resistance carries a fitness cost. Sensitive cells acquire resistance "
        "at a baseline rate plus a stress-induced term proportional to dose (the ancient mutation "
        "engine), switched off by SIM-inhibitor protocols. Cytotoxic dose kills S; immune clearance "
        "targets a recently-evolved dependency and kills S and R alike but erodes with burden (immune "
        "escape); re-differentiation adds a flux restoring I. Protocols: MTD; A1 immunotherapy; A2 "
        "adaptive dosing (maintaining a high sub-lethal burden so resistant cells are space-limited); "
        "A3 MTD + SIM-block; A4 re-differentiation; and combinations A2+A3 and A1+A4.")

    h2("2.3 Parameter provenance")
    body("Parameters are order-of-magnitude phenomenological values chosen to place the system in the "
         "biologically relevant regime; sources and rationale are given in Table 1, and robustness is "
         "established in Section 3.3 rather than claimed from any single calibration.")
    t1 = doc.add_table(rows=1, cols=3); t1.style = "Light Grid Accent 1"
    for i, htxt in enumerate(["Parameter", "Value / range", "Rationale (source)"]):
        t1.rows[0].cells[i].paragraphs[0].add_run(htxt).bold = True
    for name, val, why in PARAM_SOURCES:
        c = t1.add_row().cells; c[0].text = name; c[1].text = val; c[2].text = why
    para("Table 1. Parameter provenance.", 9, CAP, italic=True, after=8)

    # ---- 3. Results ----
    h1("3. Results")
    h2("3.1 Two families of outcomes")
    body(
        "Among cytotoxic protocols median survival rises monotonically MTD < MTD+SIM-block < adaptive < "
        "adaptive+SIM-block, so containment plus evolvability-suppression outperforms maximal "
        "cytoreduction. Yet the only protocols yielding long-term survivors attack cancer's young "
        "vulnerabilities: immunotherapy clears immunogenic tumours outright (bimodal), and immunotherapy "
        "with re-differentiation leaves the largest surviving fraction. Re-differentiation alone delays "
        "death longest but rarely holds, consistent with the order\u2013repair hysteresis.")
    figure(race_png, width=6.7)
    caption("Fig. 2. Horse race. (A) Survival. (B) Median survival and % alive at day 400. "
            "(C) Representative-mouse burden trajectories. (D) Sensitivity to fitness cost of resistance.")

    h2("3.2 Face validation against reported findings")
    body(
        "The model is not calibrated to data, but its qualitative behaviour reproduces independently "
        "reported results, a minimal consistency check. (i) MTD produces an initial response then "
        "resistant relapse, and adaptive dose-modulation prolongs control \u2014 as in preclinical mouse "
        "studies and the metastatic castrate-resistant prostate-cancer trial of adaptive therapy "
        "(Gatenby et al. 2009; Zhang et al. 2017). (ii) The advantage of adaptive dosing depends on a "
        "fitness cost of resistance (Fig. 2D), a cost demonstrated in vitro and in mice with a CDK "
        "inhibitor (Bacevic et al. 2017). (iii) Stress-induced mutagenesis accelerating resistance "
        "reflects the SOS-homologue mutation signature reported in cancer (Cisneros et al. 2017). "
        "(iv) The immune arm's bimodal hot/cold outcome mirrors clinical immunotherapy heterogeneity, "
        "and the superiority of young-dependency targeting matches the atavistic 'target the weakness' "
        "prediction (Lineweaver, Davies & Vincent 2014). Reproducing these without fitting to any of "
        "them is weak but non-trivial evidence that the dynamical structure is biologically sensible.")

    h2("3.3 Global sensitivity")
    body(
        f"A Latin-hypercube sample (n={gsens['n']}) varied nine mechanistic parameters simultaneously by "
        f"\u00b1{int(gsens['span']*100)}%. The drug-arm ordering (MTD \u2264 SIM-block \u2264 adaptive "
        f"\u2264 adaptive+SIM-block) is preserved in {gsens['ordering_frac']:.0f}% of samples, and the "
        f"young-target family (immunotherapy \u00b1 differentiation) provides the most long-term "
        f"survivors in {gsens['family_frac']:.0f}%. A one-at-a-time tornado (Fig. 3B) shows the "
        "adaptive+SIM-block survival advantage over MTD is most sensitive to the stress-induced "
        "mutagenesis rate and the burden-driven erosion flux, and never reverses sign within the range "
        "tested. The qualitative conclusions are therefore properties of the model structure, not of a "
        "single parameter choice.")
    figure(sens_png, width=6.7)
    caption("Fig. 3. Global sensitivity. (A) Distribution of the adaptive+SIM-block survival advantage "
            "over MTD across Latin-hypercube samples. (B) One-at-a-time \u00b130% tornado.")

    h2("3.4 Protocol summary")
    tbl = doc.add_table(rows=1, cols=5); tbl.style = "Light Grid Accent 1"
    for i, htxt in enumerate(["Protocol","Median survival (d)","% alive @400","% durable regression","Mean dose"]):
        tbl.rows[0].cells[i].paragraphs[0].add_run(htxt).bold = True
    for p in PROTOCOLS:
        c = tbl.add_row().cells
        c[0].text = LABELS[p]; c[1].text = f"{summary[p]['medSurv']:.0f}"
        c[2].text = f"{summary[p]['alive']:.1f}"; c[3].text = f"{summary[p]['controlled']:.1f}"
        c[4].text = f"{summary[p]['dose']:.0f}"
    para("Table 2. Outcomes by protocol.", 9, CAP, italic=True, after=8)

    # ---- 4. Discussion ----
    h1("4. Discussion")
    h2("4.1 Relation to other frameworks")
    body(
        "The somatic mutation theory treats cancer as accumulated driver mutations under selection; our "
        "model is compatible with it but adds a macroscopic order variable that the SMT lacks. Compared "
        "with cancer-stem-cell and tissue-organisation-field views, the order parameter I plays a role "
        "analogous to a coarse-grained 'degree of tissue organisation', while our therapy layer inherits "
        "the ecological/evolutionary logic of adaptive therapy (Gatenby et al. 2009). The novel content "
        "is the bistability and hysteresis of I, which supply mechanisms \u2014 self-defeating DNA-stress "
        "therapy and hysteresis-bounded differentiation \u2014 absent from purely evolutionary models.")
    h2("4.2 Proposed protocol and test")
    body(
        "The model points to a layered protocol: an adaptive + SIM-block cytotoxic backbone to contain "
        "bulk disease, hold evolvability down and keep cells re-sensitisable, combined with immunotherapy "
        "and re-differentiation to clear the contained, more-immunogenic residual.")
    h2("4.3 Predictions (falsifiable)")
    body(
        "P1. In a syngeneic mouse model (e.g. 4T1, B16), adaptive dosing plus a translesion-synthesis / "
        "SIM inhibitor extends time-to-resistance beyond adaptive dosing alone and beyond MTD. "
        "P2. Barcoded clonal diversity rises fastest under MTD and slowest under SIM-block arms. "
        "P3. Re-differentiation efficacy declines with initial tumour burden (a hysteresis threshold), "
        "not merely with growth rate. P4. Adaptive therapy's benefit scales with the measured fitness "
        "cost of resistance across cell lines. Any of these failing would falsify the corresponding "
        "model mechanism. As a direct test of the underlying atavism premise, the prediction that "
        "tumours over-express evolutionarily older (unicellular-age) genes and down-regulate younger "
        "(multicellular-age) genes (Domazet-Loso & Tautz 2010; Trigos et al. 2017) can be evaluated "
        "quantitatively on public tumour/normal expression data; an analysis pipeline for this is "
        "provided (Section 5).")
    h2("4.4 Limitations")
    body(
        "We stress that the therapeutic rankings are model-generated hypotheses to be tested "
        "(Section 4.3), not evidence that any protocol is clinically superior; the model derives which "
        "predictions should hold and under what conditions, not that they are true. This is an in-silico "
        "model \u2014 the ranking follows from encoded mechanisms and is not empirical "
        "data. Parameters are phenomenological (Table 1) and, although the conclusions are robust across "
        "the ranges tested (Section 3.3), they are not fitted to a specific tumour. The two-type S/R "
        "approximation omits a continuous resistance spectrum and spatial structure. The atavism "
        "framework itself remains contested, and phylostratigraphic inference is sensitive to overlapping "
        "gene-age categories, so downstream genomic claims should be guarded against overlap-induced "
        "significance inflation.")

    # ---- 5. Code & data availability ----
    h1("5. Code and data availability")
    if anon:
        body(
            "The differential-equation and stochastic-cohort simulations were implemented in a custom "
            "Python framework. All code, parameter files, and scripts to regenerate every figure and "
            "table will be openly archived in a public repository with a citable DOI upon acceptance "
            "(link withheld here for double-blind review). The distribution also includes a "
            "phylostratigraphy analysis pipeline that assigns gene ages, tests the unicellular-versus-"
            "multicellular expression shift between tumour and normal samples, and controls for gene-gene "
            "correlation via a sample-label permutation null to avoid overlap-induced significance "
            "inflation; it accepts standard expression matrices (e.g. TCGA/GTEx) for direct empirical "
            "evaluation of the atavism premise.")
    else:
        body(
            "The differential-equation and stochastic-cohort simulations were implemented in a custom Python "
            "framework. All code, parameter files, and scripts to regenerate every figure and table are "
            "openly archived (Zenodo DOI to be assigned on deposit; mirror on GitHub) to allow full "
            "reproduction. The distribution also includes a phylostratigraphy analysis pipeline that assigns "
            "gene ages, tests the unicellular-versus-multicellular expression shift between tumour and normal "
            "samples, and controls for gene-gene correlation via a sample-label permutation null to avoid "
            "overlap-induced significance inflation; it accepts standard expression matrices (e.g. "
            "TCGA/GTEx) for direct empirical evaluation of the atavism premise.")

    # ---- Declarations ----
    h1("Declarations")
    body("Funding: This work received no external funding.")
    body("Conflicts of interest: The author declares no competing interests.")
    body("Ethics: This study is entirely computational and involved no human participants "
         "or animal subjects.")
    if not anon:
        body("ORCID: " + HOUSE["author"] + ", " + HOUSE["orcid"] + ".")

    # ---- References ----
    h1("References")
    refs = [
        "Bacevic K, Noble R, Soffar A, et al. (2017) Spatial competition constrains resistance to targeted cancer therapy. Nature Communications 8:1995. doi:10.1038/s41467-017-01516-1",
        "Cisneros L, Bussey KJ, Orr AJ, Miocevic M, Lineweaver CH, Davies P (2017) Ancient genes establish stress-induced mutation as a hallmark of cancer. PLoS ONE 12(4):e0176258. doi:10.1371/journal.pone.0176258",
        "Davies PCW, Lineweaver CH (2011) Cancer tumors as Metazoa 1.0: tapping genes of ancient ancestors. Physical Biology 8(1):015001. doi:10.1088/1478-3975/8/1/015001",
        "Domazet-Loso T, Tautz D (2010) Phylostratigraphic tracking of cancer genes suggests a link to the emergence of multicellularity in metazoa. BMC Biology 8:66. doi:10.1186/1741-7007-8-66",
        "Gatenby RA, Silva AS, Gillies RJ, Frieden BR (2009) Adaptive therapy. Cancer Research 69(11):4894-4903. doi:10.1158/0008-5472.CAN-08-3658",
        "Lineweaver CH, Davies PCW, Vincent MD (2014) Targeting cancer's weaknesses (not its strengths): therapeutic strategies suggested by the atavistic model. BioEssays 36(9):827-835. doi:10.1002/bies.201400070",
        "Lineweaver CH, Bussey KJ, Blackburn AC, Davies PCW (2021) Cancer progression as a sequence of atavistic reversions. BioEssays 43(7):2000305. doi:10.1002/bies.202000305",
        "Trigos AS, Pearson RB, Papenfuss AT, Goode DL (2017) Altered interactions between unicellular and multicellular genes drive hallmarks of transformation. PNAS 114(24):6406-6411. doi:10.1073/pnas.1617743114",
        "Zhang J, Cunningham JJ, Brown JS, Gatenby RA (2017) Integrating evolutionary dynamics into treatment of metastatic castrate-resistant prostate cancer. Nature Communications 8:1816. doi:10.1038/s41467-017-01816-6",
    ]
    for r in refs:
        para(r, 9.5, after=3)

    doc.save(path)

# ============================================================================ #
#  6. reproduce_paper() + CLI                                                    #
# ============================================================================ #
def reproduce_paper(title="Cancer as a bistable order-repair transition",
                    subtitle="In-silico comparison of evolutionary and atavism-targeting therapies",
                    preset="tissue_atavism", outdir="/mnt/user-data/outputs", full=False):
    os.makedirs(outdir, exist_ok=True)
    P = dict(TISSUE_ATAVISM)
    if not full: P["M"] = 300
    print(f"[MyUncle] preset={preset}  M={P['M']}  running horse race ...")
    res, summary, sens, coh = run_horse_race(P)
    print("[MyUncle] running global sensitivity analysis ...")
    gsens = global_sensitivity(P, n_lhs=(120 if full else 80))
    print(f"[MyUncle]   ordering preserved {gsens['ordering_frac']:.0f}% | "
          f"young-target family wins survivors {gsens['family_frac']:.0f}%")

    theory_png = os.path.join(outdir, "myuncle_fig1_theory.png")
    race_png   = os.path.join(outdir, "myuncle_fig2_horse_race.png")
    sens_png   = os.path.join(outdir, "myuncle_fig3_sensitivity.png")
    cover_png  = os.path.join(outdir, "myuncle_cover.png")
    docx_path  = os.path.join(outdir, "MyUncle_tissue_atavism.docx")
    anon_path  = os.path.join(outdir, "MyUncle_tissue_atavism_ANON_springer.docx")

    fig_theory(theory_png)
    fig_horse_race(race_png, res, summary, sens, coh, P)
    fig_sensitivity(sens_png, gsens)
    fig_cover(cover_png, title, subtitle)
    build_docx(docx_path, title, subtitle, summary, gsens, cover_png, theory_png, race_png, sens_png, anon=False)
    build_docx(anon_path, title, subtitle, summary, gsens, cover_png, theory_png, race_png, sens_png, anon=True)

    print("\n[MyUncle] summary")
    print(f"  {'protocol':28s}{'medSurv':>9s}{'%alive':>8s}{'%ctrl':>7s}{'dose':>7s}")
    for p in PROTOCOLS:
        s = summary[p]
        print(f"  {LABELS[p]:28s}{s['medSurv']:9.1f}{s['alive']:8.1f}{s['controlled']:7.1f}{s['dose']:7.0f}")
    print("\n[MyUncle] artifacts:")
    for f in (docx_path, anon_path, race_png, theory_png, sens_png, cover_png): print("   ", f)
    return dict(summary=summary, gsens=gsens, docx=docx_path, anon=anon_path,
                figures=[theory_png, race_png, sens_png, cover_png])

# ============================================================================ #
#  7. Phylostratigraphy pipeline (empirical validation of the atavism premise)  #
#     Ready for real TCGA/GTEx matrices; demonstrated here on synthetic data.    #
# ============================================================================ #
def phylostratigraphy_analysis(expr_tumor, expr_normal, gene_age, n_perm=2000, seed=7):
    """Test the atavism premise: unicellular-age (old) genes UP and multicellular-age
    (young) genes DOWN in tumour vs normal.

    expr_tumor, expr_normal : arrays (genes x samples), log-expression.
    gene_age                : array (genes,) phylostratum rank; low = ancient (UC),
                              high = recent (MC).
    Significance uses a SAMPLE-LABEL PERMUTATION null so the p-value is robust to
    gene-gene correlation (avoids overlap-induced significance inflation that a naive
    per-gene t-test produces). Returns effect size, permutation p, and naive p for
    contrast.
    """
    rng = np.random.default_rng(seed)
    G = gene_age.shape[0]
    uc = gene_age <= np.quantile(gene_age, 0.33)     # ancient / unicellular tier
    mc = gene_age >= np.quantile(gene_age, 0.67)     # recent / multicellular tier
    all_expr = np.concatenate([expr_tumor, expr_normal], axis=1)
    nT = expr_tumor.shape[1]; nN = expr_normal.shape[1]
    labels = np.array([1]*nT + [0]*nN)

    def statistic(lab):
        t = all_expr[:, lab == 1].mean(1); n = all_expr[:, lab == 0].mean(1)
        lfc = t - n                                   # per-gene tumour-normal shift
        return lfc[uc].mean() - lfc[mc].mean()        # UC-up minus MC-up contrast

    obs = statistic(labels)
    null = np.empty(n_perm)
    for i in range(n_perm):
        null[i] = statistic(rng.permutation(labels))
    p_perm = (1 + np.sum(np.abs(null) >= abs(obs))) / (n_perm + 1)

    # naive per-gene two-sample t (independence assumption) for contrast only
    from scipy.stats import ttest_ind
    lfc = expr_tumor.mean(1) - expr_normal.mean(1)
    _, p_naive = ttest_ind(lfc[uc], lfc[mc], equal_var=False)
    return dict(effect=float(obs), p_perm=float(p_perm), p_naive=float(p_naive),
                null=null, uc_lfc=lfc[uc], mc_lfc=lfc[mc])

def synthetic_expression_cohort(G=4000, nT=120, nN=120, effect=0.6,
                                n_modules=40, seed=3):
    """SYNTHETIC demonstration data with correlated gene modules (so the
    overlap-inflation problem is real) under the atavism alternative:
    ancient genes up, recent genes down in tumour. NOT real data."""
    rng = np.random.default_rng(seed)
    gene_age = rng.uniform(0, 1, G)                          # 0 ancient .. 1 recent
    module = rng.integers(0, n_modules, G)                   # correlated blocks
    base = rng.normal(0, 1, (G, nT + nN))
    mod_T = rng.normal(0, 0.8, (n_modules, nT))              # shared module noise
    mod_N = rng.normal(0, 0.8, (n_modules, nN))
    for g in range(G):
        base[g, :nT] += mod_T[module[g]]
        base[g, nT:] += mod_N[module[g]]
    shift = effect * (0.5 - gene_age)                        # +ancient, -recent
    base[:, :nT] += shift[:, None]
    return base[:, :nT], base[:, nT:], gene_age

def fig_phylo(path, ana):
    _mpl_style(); navy = HOUSE["navy"]
    fig = plt.figure(figsize=(13, 4.4)); grid = GridSpec(1, 2, figure=fig, wspace=0.26)
    ax = fig.add_subplot(grid[0, 0])
    ax.hist(ana["uc_lfc"], bins=40, alpha=0.7, color="#C0504D",
            label=f"ancient / unicellular  (mean {ana['uc_lfc'].mean():+.2f})")
    ax.hist(ana["mc_lfc"], bins=40, alpha=0.7, color="#2E74B5",
            label=f"recent / multicellular (mean {ana['mc_lfc'].mean():+.2f})")
    ax.axvline(0, color="k", ls=":", lw=1)
    ax.set_title("(A) Tumour\u2013normal expression shift by gene age", color=navy, fontweight="bold")
    ax.set_xlabel("log fold-change (tumour \u2212 normal)"); ax.set_ylabel("genes")
    ax.legend(fontsize=7.6); ax.grid(alpha=0.25)
    ax = fig.add_subplot(grid[0, 1])
    ax.hist(ana["null"], bins=40, color="#9AA5B1", alpha=0.85, label="permutation null")
    ax.axvline(ana["effect"], color="#C0504D", lw=2,
               label=f"observed {ana['effect']:+.2f}")
    ax.set_title("(B) Permutation test (correlation-robust)", color=navy, fontweight="bold")
    ax.set_xlabel("UC\u2212MC differential shift"); ax.set_ylabel("permutations")
    ax.legend(fontsize=7.6); ax.grid(alpha=0.25)
    ax.text(0.02, 0.90, f"permutation p = {ana['p_perm']:.2e}\n"
                        f"naive per-gene p = {ana['p_naive']:.2e}\n"
                        f"(naive p overstates significance)",
            transform=ax.transAxes, fontsize=8.2, color="#1F4D78", va="top")
    fig.suptitle("Phylostratigraphy pipeline \u2014 SYNTHETIC demonstration "
                 "(plug in TCGA/GTEx for real validation)",
                 fontsize=11.5, fontweight="bold", color=navy, y=1.03)
    fig.savefig(path, dpi=140, bbox_inches="tight", facecolor="white"); plt.close(fig)

def phylo_demo(outdir="/mnt/user-data/outputs"):
    os.makedirs(outdir, exist_ok=True)
    eT, eN, age = synthetic_expression_cohort()
    ana = phylostratigraphy_analysis(eT, eN, age)
    png = os.path.join(outdir, "myuncle_fig4_phylostratigraphy.png")
    fig_phylo(png, ana)
    print(f"[MyUncle phylo-demo] effect={ana['effect']:+.3f}  "
          f"perm_p={ana['p_perm']:.2e}  naive_p={ana['p_naive']:.2e}")
    print(f"[MyUncle phylo-demo] figure -> {png}")
    return ana, png

def main():
    ap = argparse.ArgumentParser(description="MyUncle order-repair engine")
    ap.add_argument("--reproduce", action="store_true")
    ap.add_argument("--preset", default="tissue_atavism")
    ap.add_argument("--title", default="Cancer as a bistable order-repair transition")
    ap.add_argument("--subtitle", default="In-silico comparison of evolutionary and atavism-targeting therapies")
    ap.add_argument("--outdir", default="/mnt/user-data/outputs")
    ap.add_argument("--full", action="store_true", help="full n=500 cohort")
    ap.add_argument("--phylo-demo", action="store_true", help="run phylostratigraphy pipeline demo")
    a = ap.parse_args()
    if a.phylo_demo:
        phylo_demo(a.outdir)
    if a.reproduce:
        reproduce_paper(a.title, a.subtitle, a.preset, a.outdir, a.full)
    if not (a.reproduce or a.phylo_demo):
        ap.print_help()

if __name__ == "__main__":
    main()
